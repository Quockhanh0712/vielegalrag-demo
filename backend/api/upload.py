"""
Document Upload API Endpoint - File Processing and Indexing.
"""
import uuid
import os
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete

from backend.config import settings, DATA_DIR
from backend.db.database import get_db
from backend.db.models import UserDocument
from backend.utils.logger import get_logger
from backend.utils.exceptions import LegalRAGException

logger = get_logger("api.upload")

router = APIRouter()

# Supported file types
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

# Upload directory
UPLOAD_DIR = DATA_DIR / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def validate_file(file: UploadFile) -> None:
    """Validate uploaded file."""
    # Check extension
    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type '{ext}' not supported. Allowed: {ALLOWED_EXTENSIONS}"
        )


async def extract_text(file_path: Path, extension: str) -> str:
    """Extract text from file with multiple fallback methods."""
    text = ""
    
    try:
        if extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
                
        elif extension == ".pdf":
            text = await _extract_pdf_text(file_path)
            
        elif extension == ".docx":
            from docx import Document
            doc = Document(file_path)
            text = "\n".join(para.text for para in doc.paragraphs)
            
    except Exception as e:
        logger.error(f"Failed to extract text: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to extract text: {str(e)}")
    
    return text


async def _extract_pdf_text(file_path: Path) -> str:
    """Extract text from PDF with 3 fallback methods."""
    errors = []
    
    # Method 1: pypdf (most reliable for simple PDFs)
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if text.strip():
            logger.info("PDF extracted with pypdf")
            return text
    except ImportError:
        errors.append("pypdf not installed")
    except Exception as e:
        errors.append(f"pypdf failed: {str(e)[:50]}")
    
    # Method 2: pypdfium2 (better for complex layouts)
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(str(file_path))
        text_parts = []
        for page in pdf:
            textpage = page.get_textpage()
            text_parts.append(textpage.get_text_range())
        pdf.close()
        text = "\n".join(text_parts)
        if text.strip():
            logger.info("PDF extracted with pypdfium2")
            return text
    except ImportError:
        errors.append("pypdfium2 not installed")
    except Exception as e:
        errors.append(f"pypdfium2 failed: {str(e)[:50]}")
    
    # All methods failed
    error_msg = " | ".join(errors)
    raise HTTPException(
        status_code=400, 
        detail=f"PDF không đọc được. Errors: {error_msg}"
    )


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into chunks with overlap, recognizing legal structure."""
    if not text:
        return []
    
    # Try legal-specific chunking first
    legal_chunks = _chunk_legal_text(text)
    if legal_chunks:
        return legal_chunks
    
    # Fallback to standard chunking
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to break at sentence boundary
        if end < len(text):
            for sep in [". ", ".\n", "\n\n", "\n"]:
                pos = text.rfind(sep, start, end)
                if pos > start + chunk_size // 2:
                    end = pos + len(sep)
                    break
        
        chunks.append(text[start:end].strip())
        start = end - overlap
    
    return [c for c in chunks if c]  # Remove empty chunks


def _chunk_legal_text(text: str) -> List[str]:
    """Chunk legal text by Điều/Khoản structure."""
    import re
    
    # Pattern for Vietnamese legal articles: "Điều 1.", "Điều 53", etc.
    dieu_pattern = r'(Điều\s+\d+[a-z]?\.?)'
    
    # Split by Điều
    parts = re.split(dieu_pattern, text, flags=re.IGNORECASE)
    
    if len(parts) < 3:  # No Điều found
        return []
    
    chunks = []
    current_dieu = ""
    current_text = ""
    
    for i, part in enumerate(parts):
        if re.match(dieu_pattern, part, re.IGNORECASE):
            # Save previous Điều if exists
            if current_dieu and current_text.strip():
                chunks.append(f"{current_dieu}\n{current_text.strip()}")
            current_dieu = part.strip()
            current_text = ""
        else:
            current_text += part
    
    # Don't forget the last Điều
    if current_dieu and current_text.strip():
        chunks.append(f"{current_dieu}\n{current_text.strip()}")
    
    logger.info(f"Legal chunking: {len(chunks)} Điều extracted")
    return chunks


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    user_id: str = Form(...),
    session_id: str = Form(...),
    db: AsyncSession = Depends(get_db)
):
    """
    Upload and index a document for user-specific search.
    
    Flow:
    1. Validate file
    2. Save to disk
    3. Extract text
    4. Chunk text
    5. Embed chunks
    6. Store in Qdrant (user_docs_private collection)
    7. Save metadata to database
    
    Args:
        file: Uploaded file (PDF, DOCX, TXT)
        user_id: User identifier
        session_id: Session identifier
        
    Returns:
        Upload status and document info
    """
    logger.info(f"Upload request: file={file.filename}, user={user_id}")
    
    # Validate
    validate_file(file)
    
    # Generate document ID
    doc_id = f"user_{user_id}_{uuid.uuid4().hex[:8]}"
    extension = Path(file.filename).suffix.lower()
    
    # Save file
    file_path = UPLOAD_DIR / f"{doc_id}{extension}"
    
    try:
        # Read and save file
        content = await file.read()
        
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail=f"File too large. Max: {MAX_FILE_SIZE // 1024 // 1024}MB")
        
        with open(file_path, "wb") as f:
            f.write(content)
        
        # Extract text
        text = await extract_text(file_path, extension)
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No text content found in file")
        
        # Chunk text
        chunks = chunk_text(text)
        logger.info(f"Document chunked: {len(chunks)} chunks")
        
        # Embed and store
        from backend.core.embeddings import get_embedding_model
        from backend.core.qdrant_store import get_qdrant_connector
        
        embedding_model = get_embedding_model()
        qdrant = get_qdrant_connector()
        
        # Embed all chunks
        embeddings = embedding_model.embed(chunks, show_progress=False)
        
        # Prepare points for Qdrant
        points = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            points.append({
                "id": f"{doc_id}_chunk_{i}",
                "vector": embedding.tolist(),
                "payload": {
                    "text": chunk,
                    "doc_id": doc_id,
                    "user_id": user_id,
                    "session_id": session_id,
                    "file_name": file.filename,
                    "chunk_index": i,
                    "source_type": "user_document"
                }
            })
        
        # Insert into Qdrant
        success = qdrant.insert_points(
            collection_name=settings.QDRANT_USER_COLLECTION,
            points=points
        )
        
        if not success:
            raise HTTPException(status_code=500, detail="Failed to store document vectors")
        
        # Save to database
        user_doc = UserDocument(
            user_id=user_id,
            session_id=session_id,
            doc_id=doc_id,
            file_name=file.filename,
            file_size=len(content),
            num_chunks=len(chunks),
            upload_status="completed",
            created_at=datetime.utcnow()
        )
        db.add(user_doc)
        await db.commit()
        
        logger.info(f"Document uploaded: doc_id={doc_id}, chunks={len(chunks)}")
        
        return {
            "status": "success",
            "doc_id": doc_id,
            "file_name": file.filename,
            "file_size": len(content),
            "num_chunks": len(chunks),
            "message": f"Document indexed with {len(chunks)} chunks"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"Upload failed: {e}")
        # Cleanup on failure
        if file_path.exists():
            file_path.unlink()
        await db.rollback()
        raise HTTPException(status_code=500, detail={"error": "UploadFailed", "message": str(e)})


@router.get("/documents/{user_id}")
async def list_documents(
    user_id: str,
    db: AsyncSession = Depends(get_db)
):
    """List all documents uploaded by a user."""
    try:
        result = await db.execute(
            select(UserDocument)
            .where(UserDocument.user_id == user_id)
            .order_by(UserDocument.created_at.desc())
        )
        documents = result.scalars().all()
        
        return {
            "documents": [
                {
                    "doc_id": doc.doc_id,
                    "file_name": doc.file_name,
                    "file_size": doc.file_size,
                    "num_chunks": doc.num_chunks,
                    "status": doc.upload_status,
                    "created_at": doc.created_at.isoformat()
                }
                for doc in documents
            ],
            "total": len(documents)
        }
        
    except Exception as e:
        logger.error(f"Failed to list documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/documents/{doc_id}")
async def delete_document(
    doc_id: str,
    user_id: str,
    db: AsyncSession = Depends(get_db)
):
    """Delete a user document from both Qdrant and database."""
    try:
        # Verify ownership
        result = await db.execute(
            select(UserDocument).where(
                UserDocument.doc_id == doc_id,
                UserDocument.user_id == user_id
            )
        )
        doc = result.scalar_one_or_none()
        
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Delete from Qdrant
        from backend.core.qdrant_store import get_qdrant_connector
        qdrant = get_qdrant_connector()
        
        qdrant.delete_by_filter(
            collection_name=settings.QDRANT_USER_COLLECTION,
            filter_conditions={"doc_id": doc_id}
        )
        
        # Delete from database
        await db.delete(doc)
        await db.commit()
        
        # Delete file if exists
        for ext in ALLOWED_EXTENSIONS:
            file_path = UPLOAD_DIR / f"{doc_id}{ext}"
            if file_path.exists():
                file_path.unlink()
                break
        
        logger.info(f"Document deleted: {doc_id}")
        
        return {"status": "deleted", "doc_id": doc_id}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete document: {e}")
        await db.rollback()
        raise HTTPException(status_code=500, detail=str(e))
